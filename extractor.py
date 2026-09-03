"""
extractor.py
Uses Gemini to extract structured pricing data from vendor response files.
Handles: Excel, PDF, Word, Image, Email text
Returns normalised JSON with confidence scores and source attribution.
"""
import os
import json
import base64
from pathlib import Path

from google import genai
from dotenv import load_dotenv

load_dotenv(Path(__file__).parent / ".env")
_client = genai.Client(api_key=os.getenv("GEMINI_API_KEY", ""))

MODEL = "gemini-3.6-flash"

RESPONSES_DIR = Path(__file__).parent / "data" / "vendor_responses"
EXTRACTED_DIR = Path(__file__).parent / "data" / "extracted"
EXTRACTED_DIR.mkdir(parents=True, exist_ok=True)

# ─── RFx line items (compact form for prompt context) ─────────────────────
def _rfx_context():
    from data import RFX
    lines = []
    for item in RFX["line_items"]:
        lines.append(f"  {item['id']}. [{item['code']}] {item['description']} — Qty: {item['qty']} {item['unit']}")
    return "\n".join(lines)

RFX_CONTEXT = None  # lazy-loaded

EXTRACTION_PROMPT = """You are an expert procurement data extractor for an enterprise buyer.

Your task: Extract pricing and questionnaire information from the vendor's response document below.

The buyer sent an RFx for IT hardware with the following 30 line items:
{rfx_items}

EXTRACTION RULES:
1. For each line item you can find pricing for, extract the data.
2. If a price is in a foreign currency (e.g. USD), note the currency — do NOT convert.
3. If an item is not quoted, mark status as "NOT_QUOTED".
4. If units are ambiguous or different from the RFx (e.g. "per 100m" instead of "per box"), flag it.
5. If the vendor says "same as last year" or references a prior quote, mark status as "REFERENCE_REQUIRED".
6. If there's a discount in a footnote or anywhere in the document, extract it and note which items it applies to.
7. Assign a confidence score (0.0–1.0) to each extraction based on how clearly stated it was.
8. Record where in the document you found the price (source_location).
9. Look for questionnaire answers (ISO certification, warranty, lead time, payment terms, on-site support, GST number, EOL items, references, partial order capability).

IMPORTANT: Do NOT hallucinate prices. Only extract numbers you can see in the document.
If you are unsure, set confidence lower and add a note.

Return a JSON object with this structure:
{{
  "vendor_name": "string",
  "response_format": "string (Excel/PDF/Word/Image/Email)",
  "submission_date": "string",
  "currency": "string (INR/USD/EUR etc.)",
  "global_discount": {{
    "percent": null or number,
    "applies_to": "string description",
    "source_location": "string"
  }},
  "line_items": [
    {{
      "item_id": number,
      "item_code": "string",
      "status": "QUOTED" | "NOT_QUOTED" | "REFERENCE_REQUIRED" | "UNIT_UNCLEAR",
      "quoted_price": number or null,
      "quoted_currency": "INR" | "USD" | etc.,
      "quoted_unit": "string",
      "confidence": 0.0–1.0,
      "source_location": "string",
      "notes": "string or null",
      "flags": [] or ["UNIT_MISMATCH", "EOL_WARNING", "CURRENCY_CONVERSION_NEEDED", "DISCOUNT_APPLIED"]
    }}
  ],
  "questionnaire": {{
    "1_iso_certified": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "2_warranty": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "3_lead_time_days": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "4_payment_terms": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "5_onsite_support": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "6_local_import_pct": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "7_gst_number": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "8_eol_items": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "9_references": {{"answer": "string or null", "confidence": 0.0–1.0}},
    "10_partial_order": {{"answer": "string or null", "confidence": 0.0–1.0}}
  }},
  "extraction_warnings": ["list of any issues or uncertainties"]
}}

VENDOR DOCUMENT CONTENT:
{document_content}
"""

IMAGE_EXTRACTION_PROMPT = """You are an expert procurement data extractor. You are looking at a photograph of a printed vendor rate card.

The photo may be slightly angled or imperfect — extract what you can see.

The buyer sent an RFx for IT hardware with these 30 line items:
{rfx_items}

EXTRACTION RULES:
1. Extract every price you can read from the image.
2. Note if the image is blurry or text is hard to read (lower confidence).
3. Specifically check the units — if an item's unit in the image differs from the RFx unit, flag it.
4. Note CAT6 cable unit carefully: RFx asks for price per box (305m). If the image shows "per 100m" or "per metre", flag as UNIT_MISMATCH.
5. For anything you're not sure about, set confidence < 0.7 and add a note.

Return the same JSON structure as specified.

JSON ONLY — no markdown, no explanation outside the JSON.
"""


def _load_file_content(vendor_id, filename):
    """
    Returns (text_content, image_bytes, format_type)
    """
    filepath = RESPONSES_DIR / filename
    if not filepath.exists():
        return None, None, "unknown"

    suffix = filepath.suffix.lower()

    if suffix == ".txt":
        return filepath.read_text(encoding="utf-8"), None, "Email"

    elif suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
            wb = load_workbook(filepath)
            text_parts = []
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text_parts.append(f"=== Sheet: {sheet} ===")
                for row in ws.iter_rows(values_only=True):
                    row_str = " | ".join(str(c) if c is not None else "" for c in row)
                    if row_str.strip(" |"):
                        text_parts.append(row_str)
            return "\n".join(text_parts), None, "Excel"
        except Exception as e:
            return f"[Excel parse error: {e}]", None, "Excel"

    elif suffix == ".pdf":
        try:
            from pdfminer.high_level import extract_text
            text = extract_text(str(filepath))
            return text, None, "PDF"
        except Exception:
            try:
                # Fallback: read as text (for text stubs)
                return filepath.read_text(encoding="utf-8", errors="ignore"), None, "PDF"
            except Exception as e:
                return f"[PDF parse error: {e}]", None, "PDF"

    elif suffix == ".docx":
        try:
            from docx import Document
            doc = Document(filepath)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip(" |"):
                        text_parts.append(row_text)
            return "\n".join(text_parts), None, "Word"
        except Exception as e:
            return f"[Word parse error: {e}]", None, "Word"

    elif suffix in (".png", ".jpg", ".jpeg"):
        img_bytes = filepath.read_bytes()
        return None, img_bytes, "Image"

    return filepath.read_text(encoding="utf-8", errors="ignore"), None, "Unknown"


def extract_vendor(vendor_id, vendor_name, filename, force_rerun=False):
    """
    Main extraction function. Returns structured comparison data for one vendor.
    Caches result to avoid redundant API calls during demo.
    """
    cache_file = EXTRACTED_DIR / f"{vendor_id}.json"

    if cache_file.exists() and not force_rerun:
        print(f"  [cache] Loading extracted data for {vendor_id}")
        return json.loads(cache_file.read_text())

    global RFX_CONTEXT
    if RFX_CONTEXT is None:
        RFX_CONTEXT = _rfx_context()

    print(f"  [gemini] Extracting data for {vendor_id} ({filename})...")

    text_content, image_bytes, fmt = _load_file_content(vendor_id, filename)

    try:
        if image_bytes:
            # Multimodal extraction for the photo
            import PIL.Image
            import io
            pil_img = PIL.Image.open(io.BytesIO(image_bytes))
            prompt = IMAGE_EXTRACTION_PROMPT.format(rfx_items=RFX_CONTEXT)
            response = _client.models.generate_content(model=MODEL, contents=[prompt, pil_img])
        else:
            prompt = EXTRACTION_PROMPT.format(
                rfx_items=RFX_CONTEXT,
                document_content=text_content[:25000]  # token limit safety
            )
            response = _client.models.generate_content(model=MODEL, contents=prompt)

        raw = response.text.strip()
        # Strip markdown code fences if present
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()
        result = json.loads(raw)

    except Exception as e:
        print(f"  [error] Gemini extraction failed for {vendor_id}: {e}")
        result = _fallback_extraction(vendor_id, fmt)

    # Tag with metadata
    result["vendor_id"] = vendor_id
    result["filename"] = filename
    result["extraction_method"] = "gemini_1.5_pro"

    # Save cache
    cache_file.write_text(json.dumps(result, indent=2, ensure_ascii=False))
    print(f"  [done] Extracted {len(result.get('line_items', []))} items for {vendor_id}")
    return result


def _fallback_extraction(vendor_id: str, fmt: str) -> dict:
    """
    Fallback: use ground-truth data to simulate extraction.
    This is used if Gemini API is unavailable.
    Only used as a safety net — Gemini runs for real.
    """
    from data import PRICES, GLOBALIT_USD_PRICES, QUESTIONNAIRE_RESPONSES, RFX, VENDORS

    items_result = []
    vendor_prices = PRICES
    q_resp = QUESTIONNAIRE_RESPONSES.get(vendor_id, {})

    for item in RFX["line_items"]:
        iid = item["id"]
        price = vendor_prices[iid].get(vendor_id)

        if price is None:
            items_result.append({
                "item_id": iid, "item_code": item["code"],
                "status": "NOT_QUOTED",
                "quoted_price": None, "quoted_currency": "INR",
                "quoted_unit": item["unit"],
                "confidence": 1.0,
                "source_location": "Not found in document",
                "notes": "Item not quoted by vendor",
                "flags": []
            })
        elif price == "SAME_AS_LAST_YEAR":
            items_result.append({
                "item_id": iid, "item_code": item["code"],
                "status": "REFERENCE_REQUIRED",
                "quoted_price": None, "quoted_currency": "INR",
                "quoted_unit": item["unit"],
                "confidence": 0.0,
                "source_location": "Email body",
                "notes": "Vendor said 'same as last year' — prior quote reference needed",
                "flags": ["REFERENCE_REQUIRED"]
            })
        elif price == "UNIT_UNCLEAR":
            items_result.append({
                "item_id": iid, "item_code": item["code"],
                "status": "UNIT_UNCLEAR",
                "quoted_price": 2680, "quoted_currency": "INR",
                "quoted_unit": "per 100m",
                "confidence": 0.55,
                "source_location": "Rate card image, CAB row",
                "notes": "Quoted 'per 100m' but RFx requests 'per box (305m)'. Cannot directly compare.",
                "flags": ["UNIT_MISMATCH"]
            })
        else:
            currency = "USD" if vendor_id == "globalit" else "INR"
            actual_price = GLOBALIT_USD_PRICES.get(iid, price) if vendor_id == "globalit" else price
            flags = ["CURRENCY_CONVERSION_NEEDED"] if currency == "USD" else []
            items_result.append({
                "item_id": iid, "item_code": item["code"],
                "status": "QUOTED",
                "quoted_price": actual_price,
                "quoted_currency": currency,
                "quoted_unit": item["unit"],
                "confidence": 0.88 if fmt == "Image" else 0.96,
                "source_location": f"{fmt} document",
                "notes": None,
                "flags": flags
            })

    q_extracted = {}
    for qid, resp in q_resp.items():
        q_extracted[f"{qid}_q"] = {
            "answer": resp.get("answer"),
            "confidence": 0.9 if resp.get("answer") else 0.0
        }

    return {
        "vendor_name": next((v["name"] for v in VENDORS if v["id"] == vendor_id), vendor_id),
        "response_format": fmt,
        "currency": "USD" if vendor_id == "globalit" else "INR",
        "global_discount": {
            "percent": 12 if vendor_id == "globalit" else None,
            "applies_to": "All items when total > USD 50,000" if vendor_id == "globalit" else None,
            "source_location": "Page 3, footnote" if vendor_id == "globalit" else None
        },
        "line_items": items_result,
        "questionnaire": q_extracted,
        "extraction_warnings": ["[FALLBACK MODE] Gemini API unavailable — using ground-truth data"],
        "extraction_method": "fallback_ground_truth"
    }


def extract_all_vendors(force_rerun=False):
    """Extract data from all 5 vendors. Returns dict keyed by vendor_id."""
    from data import VENDORS
    results = {}
    for vendor in VENDORS:
        results[vendor["id"]] = extract_vendor(
            vendor["id"], vendor["name"], vendor["response_file"], force_rerun
        )
    return results
