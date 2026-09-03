"""
create_vendor_files.py
Generates 5 realistic vendor response files in different formats:
  1. TechPro Solutions   → Excel (.xlsx)
  2. GlobalIT Supplies   → PDF
  3. QuickByte India     → Word (.docx)
  4. DigitalEdge Corp    → Image (angled phone photo)
  5. Shree IT Traders    → Plain email text (.txt)
"""
import os
import sys
from pathlib import Path

# ── Paths ────────────────────────────────────────────────────────────────────
BASE = Path(__file__).parent
RESPONSES_DIR = BASE / "data" / "vendor_responses"
RESPONSES_DIR.mkdir(parents=True, exist_ok=True)

sys.path.insert(0, str(BASE))
from data import RFX, VENDORS, PRICES, GLOBALIT_USD_PRICES, QUESTIONNAIRE_RESPONSES

ITEMS = RFX["line_items"]


# ════════════════════════════════════════════════════════════════════════════
# 1. TechPro — Excel
# ════════════════════════════════════════════════════════════════════════════
def create_techpro_excel():
    from openpyxl import Workbook
    from openpyxl.styles import (PatternFill, Font, Alignment, Border, Side,
                                  numbers)
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws_price = wb.active
    ws_price.title = "Commercial Bid"
    ws_q = wb.create_sheet("Questionnaire Responses")

    # Header styling
    header_fill = PatternFill("solid", fgColor="003366")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    alt_fill    = PatternFill("solid", fgColor="EBF3FB")

    # ── Commercial Bid sheet ────────────────────────────────────────────────
    headers = ["#", "Item Code", "Description", "Category", "Qty", "Unit",
               "Unit Price (INR)", "Total (INR)", "Brand / Make", "Warranty", "Notes"]
    ws_price.append(["TechPro Solutions Pvt. Ltd. — Commercial Bid"])
    ws_price.append(["RFx Ref: RFX-2024-ITH-001", "Date: 22-Nov-2024", "", "", "", "", "", "", "", "", ""])
    ws_price.append([])
    ws_price.append(headers)

    # Style header row (row 4)
    for col_idx, h in enumerate(headers, 1):
        cell = ws_price.cell(row=4, column=col_idx)
        cell.fill   = header_fill
        cell.font   = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    missing = {3, 7, 19}  # TechPro misses Blade Chassis, CAD Workstation, UPS 10kVA

    brands = {
        "Servers": "Dell PowerEdge", "Laptops": "HP EliteBook", "Workstations": "HP Z8",
        "Desktops": "Lenovo ThinkCentre", "Networking": "Cisco Catalyst",
        "Power": "APC by Schneider", "Infrastructure": "Vertiv / generic",
        "Cabling": "Schneider / Molex", "Storage": "Synology / Seagate",
        "Components": "Samsung / Micron", "Peripherals": "Dell UltraSharp",
    }

    row = 5
    grand_total = 0
    for item in ITEMS:
        iid = item["id"]
        if iid in missing:
            continue
        price = PRICES[iid]["techpro"]
        total = price * item["qty"]
        grand_total += total
        data = [
            iid, item["code"], item["description"], item["category"],
            item["qty"], item["unit"],
            price, total,
            brands.get(item["category"], "OEM"),
            "36 months on-site" if item["category"] == "Servers" else "12 months",
            ""
        ]
        ws_price.append(data)
        if row % 2 == 0:
            for c in range(1, 12):
                ws_price.cell(row=row, column=c).fill = alt_fill
        # Format price cols
        ws_price.cell(row=row, column=7).number_format = '₹#,##0'
        ws_price.cell(row=row, column=8).number_format = '₹#,##0'
        row += 1

    # Grand total row
    ws_price.append(["", "", "", "", "", "GRAND TOTAL", "", grand_total, "", "", ""])
    ws_price.cell(row=row, column=8).number_format = '₹#,##0'
    ws_price.cell(row=row, column=6).font = Font(bold=True)
    ws_price.cell(row=row, column=8).font = Font(bold=True)

    # Note about missing items
    ws_price.append([])
    ws_price.append(["NOTE: Items SRV-003 (Blade Server Chassis), WKS-001 (CAD Workstation), "
                     "and UPS-003 (10 kVA UPS) are not quoted. "
                     "Contact sales@techpro-solutions.in for custom configurations."])

    # Column widths
    widths = [4, 10, 35, 15, 6, 8, 16, 16, 20, 22, 20]
    for i, w in enumerate(widths, 1):
        ws_price.column_dimensions[get_column_letter(i)].width = w

    # ── Questionnaire sheet ─────────────────────────────────────────────────
    q_resp = QUESTIONNAIRE_RESPONSES["techpro"]
    ws_q.append(["TechPro Solutions — Questionnaire Responses"])
    ws_q.append([])
    qs = RFX["questionnaire"]
    for q in qs:
        r = q_resp[q["id"]]
        ws_q.append([f"Q{q['id']}. {q['question']}"])
        ws_q.append([f"Answer: {r['answer']}"])
        if r.get("detail"):
            ws_q.append([f"       {r['detail']}"])
        ws_q.append([])

    ws_q.column_dimensions["A"].width = 90

    out = RESPONSES_DIR / "techpro_response.xlsx"
    wb.save(out)
    print(f"✓ Created {out}")


# ════════════════════════════════════════════════════════════════════════════
# 2. GlobalIT — PDF
# ════════════════════════════════════════════════════════════════════════════
def create_globalit_pdf():
    try:
        from fpdf import FPDF
    except ImportError:
        print("  fpdf2 not installed — creating text stub for GlobalIT PDF")
        _create_globalit_text_stub()
        return

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 12)
            self.set_fill_color(0, 82, 155)
            self.set_text_color(255, 255, 255)
            self.cell(0, 10, "GlobalIT Supplies Inc. — Request for Quote Response", fill=True, ln=True, align="C")
            self.set_text_color(0, 0, 0)
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 10, f"Page {self.page_no()} | GlobalIT Supplies Inc. | CONFIDENTIAL", align="C")

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    pdf.set_font("Helvetica", "", 10)

    # Cover info
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 7, "Quotation Details", ln=True)
    pdf.set_font("Helvetica", "", 10)
    info = [
        ("Ref No:", "GI-Q-2024-4487"),
        ("RFx Ref:", "RFX-2024-ITH-001"),
        ("Date:", "23 November 2024"),
        ("Valid Until:", "22 January 2025 (60 days)"),
        ("Contact:", "bids@globalit-supplies.com | +1-800-GLOBALIT"),
        ("Currency:", "USD (United States Dollar)"),
    ]
    for label, val in info:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(40, 6, label)
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, val, ln=True)

    pdf.ln(4)
    pdf.set_font("Helvetica", "", 9)
    pdf.set_fill_color(255, 240, 200)
    pdf.multi_cell(0, 6,
        "IMPORTANT: All prices are quoted in USD. Buyer to convert to INR at applicable exchange rate on date of PO. "
        "Please refer to footnote on Page 3 regarding applicable volume discount.", fill=True)
    pdf.ln(4)

    # Pricing table
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(0, 82, 155)
    pdf.set_text_color(255, 255, 255)
    col_w = [8, 14, 60, 12, 12, 22, 28]
    headers = ["#", "Code", "Description", "Qty", "Unit", "Unit Price (USD)", "Total (USD)"]
    for w, h in zip(col_w, headers):
        pdf.cell(w, 7, h, border=1, fill=True, align="C")
    pdf.ln()
    pdf.set_text_color(0, 0, 0)

    missing_globalit = {7, 19, 28}
    grand_usd = 0
    fill_toggle = False
    for item in ITEMS:
        iid = item["id"]
        if iid in missing_globalit:
            continue
        usd_price = GLOBALIT_USD_PRICES.get(iid, 0)
        total_usd = usd_price * item["qty"]
        grand_usd += total_usd
        pdf.set_font("Helvetica", "", 8)
        if fill_toggle:
            pdf.set_fill_color(235, 243, 251)
        else:
            pdf.set_fill_color(255, 255, 255)
        row_data = [
            str(iid), item["code"], item["description"][:45],
            str(item["qty"]), item["unit"],
            f"${usd_price:,.2f}", f"${total_usd:,.2f}"
        ]
        for w, d in zip(col_w, row_data):
            pdf.cell(w, 6, d, border=1, fill=True, align="C" if d[0].isdigit() or d[0]=="$" else "L")
        pdf.ln()
        fill_toggle = not fill_toggle

    # Total
    pdf.set_font("Helvetica", "B", 10)
    pdf.set_fill_color(220, 230, 241)
    pdf.cell(sum(col_w[:-1]), 7, "GRAND TOTAL (before discount — see footnote Pg. 3)", border=1, fill=True, align="R")
    pdf.cell(col_w[-1], 7, f"${grand_usd:,.2f}", border=1, fill=True, align="C")
    pdf.ln()

    # Items not quoted
    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(150, 0, 0)
    pdf.multi_cell(0, 5,
        "Items not quoted: WKS-001 (CAD Workstation) — not in product portfolio. "
        "UPS-003 (10 kVA) — lead time exceeds RFx delivery window. "
        "GPU-001 (Data Center GPU) — unavailable at time of quoting.")
    pdf.set_text_color(0, 0, 0)

    # Questionnaire page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Section 2: Questionnaire Responses", ln=True)
    pdf.set_font("Helvetica", "", 9)
    q_resp = QUESTIONNAIRE_RESPONSES["globalit"]
    for q in RFX["questionnaire"]:
        r = q_resp[q["id"]]
        pdf.set_font("Helvetica", "B", 9)
        pdf.multi_cell(0, 5, f"Q{q['id']}. {q['question']}")
        pdf.set_font("Helvetica", "", 9)
        ans = r["answer"] or "Not answered"
        pdf.multi_cell(0, 5, f"   {ans}")
        if r.get("flag") == "EOL_WARNING":
            pdf.set_fill_color(255, 200, 200)
            pdf.set_font("Helvetica", "B", 8)
            pdf.multi_cell(0, 5, "   ⚠ EOL NOTICE: LPT-002 reaches end-of-life in Q2 2026.", fill=True)
            pdf.set_font("Helvetica", "", 9)
        pdf.ln(2)

    # Footnote page — BURIED DISCOUNT
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(0, 8, "Section 3: Terms & Conditions", ln=True)
    pdf.set_font("Helvetica", "", 9)
    lorem = (
        "All products are supplied by GlobalIT Supplies Inc., an authorised distributor for "
        "major OEM brands. Products are subject to availability. Lead times quoted assume "
        "standard stock availability. Custom configurations may extend lead times. "
        "GlobalIT Supplies Inc. complies with all applicable export control regulations. "
        "Prices are EXW Singapore; freight and customs duties are to be borne by the buyer. "
        "Warranty claims to be initiated through GlobalIT Supplies India liaison office. "
        "Force majeure clauses apply as per standard international trade terms.\n\n"
        "Payment: Strictly Net 45 days from date of invoice. Late payments attract 1.5% per month.\n\n"
        "Dispute Resolution: All disputes under this quotation are subject to Singapore International "
        "Arbitration Centre (SIAC) rules. Governing law: Republic of Singapore.\n\n"
        "Limitation of Liability: GlobalIT Supplies shall not be liable for indirect, consequential, "
        "or punitive damages. Maximum liability limited to invoice value of affected items.\n\n"
    )
    pdf.multi_cell(0, 5, lorem)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(80, 80, 80)
    # The buried footnote
    pdf.multi_cell(0, 4,
        "* Volume Discount: A 12% discount applies to all line items in this quotation when the "
        "total order value (net) exceeds USD 50,000. The above-stated unit prices are list prices; "
        "the final invoice will reflect the 12% discount automatically. Buyer should use discounted "
        "prices for budgeting: multiply listed unit price by 0.88. "
        "This discount is subject to payment within Net 45 days; late payments forfeit the discount.")

    out = RESPONSES_DIR / "globalit_response.pdf"
    pdf.output(str(out))
    print(f"✓ Created {out}")


def _create_globalit_text_stub():
    """Fallback: create a text file that looks like PDF-extracted content."""
    out = RESPONSES_DIR / "globalit_response.txt"
    lines = ["GlobalIT Supplies Inc. — RFx Response (RFX-2024-ITH-001)",
             "Date: 23 November 2024 | Currency: USD",
             "=" * 60, ""]
    missing_globalit = {7, 19, 28}
    for item in ITEMS:
        iid = item["id"]
        if iid in missing_globalit:
            continue
        usd = GLOBALIT_USD_PRICES.get(iid, 0)
        lines.append(f"{item['code']} | {item['description']} | {item['qty']} {item['unit']} | ${usd}/unit")
    lines += ["", "NOTE (Pg 3 footnote): 12% volume discount applies to all items when total > USD 50,000.",
              "LPT-002 is scheduled EOL in Q2 2026."]
    out.write_text("\n".join(lines))
    print(f"✓ Created text stub: {out}")


# ════════════════════════════════════════════════════════════════════════════
# 3. QuickByte — Word doc (prices buried in paragraphs)
# ════════════════════════════════════════════════════════════════════════════
def create_quickbyte_docx():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("QuickByte India — Commercial Quotation", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Reference: RFX-2024-ITH-001 | Date: 24 November 2024")
    doc.add_paragraph("Dear Procurement Team,")
    doc.add_paragraph(
        "Thank you for the opportunity to participate in your IT hardware procurement. "
        "QuickByte India is pleased to submit our most competitive pricing for the items listed below. "
        "We pride ourselves on offering the best value in the market with fast delivery from our "
        "Bengaluru and Hyderabad warehouses."
    )

    # Prose-style pricing — the edge case
    missing_qb = {3, 14, 16, 19, 28}
    categories = {}
    for item in ITEMS:
        if item["id"] not in missing_qb:
            categories.setdefault(item["category"], []).append(item)

    for cat, items in categories.items():
        doc.add_heading(cat, level=1)
        for item in items:
            iid = item["id"]
            price = PRICES[iid]["quickbyte"]
            total = price * item["qty"]
            # Write as paragraph prose
            p = doc.add_paragraph()
            p.add_run(f"{item['description']} ({item['code']}): ").bold = True
            p.add_run(
                f"We can supply {item['qty']} {item['unit']}(s) at a unit price of "
                f"Rs. {price:,}/- (Rupees {_inr_words(price)} only). "
                f"Total value: Rs. {total:,}/-. "
                f"Brand: {'Dell/HP' if cat in ('Servers','Laptops','Workstations','Desktops') else 'Cisco/Aruba' if cat=='Networking' else 'APC/Vertiv' if cat=='Power' else 'Generic/OEM'}. "
                f"Specs as per RFx. Delivery in 14 days."
            )

    # Note items not quoted
    doc.add_heading("Items Not Quoted", level=1)
    doc.add_paragraph(
        "We regret that we are unable to quote the following items at this time: "
        "SRV-003 (Blade Server Chassis) — not in our portfolio; "
        "FWL-001 (Enterprise Firewall) — procurement lead time exceeds 30 days; "
        "WLC-001 (Wireless Controller) — out of stock; "
        "UPS-003 (10 kVA UPS) — not available; "
        "GPU-001 (Data Center GPU) — allocation constraints. "
        "Please do consider us for the remaining items."
    )

    # Questionnaire
    doc.add_heading("Questionnaire Responses", level=1)
    q_resp = QUESTIONNAIRE_RESPONSES["quickbyte"]
    for q in RFX["questionnaire"]:
        r = q_resp[q["id"]]
        doc.add_heading(f"Q{q['id']}. {q['question']}", level=2)
        doc.add_paragraph(r["answer"] or "Not answered")

    # Summary table at the end
    doc.add_heading("Summary Pricing Table", level=1)
    doc.add_paragraph("(For convenience — binding prices are stated in the text above)")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Code", "Description", "Qty", "Unit Price (₹)", "Total (₹)"]):
        hdr[i].text = h

    for item in ITEMS:
        iid = item["id"]
        if iid in missing_qb:
            continue
        price = PRICES[iid]["quickbyte"]
        row = table.add_row().cells
        row[0].text = item["code"]
        row[1].text = item["description"]
        row[2].text = str(item["qty"])
        row[3].text = f"₹{price:,}"
        row[4].text = f"₹{price * item['qty']:,}"

    out = RESPONSES_DIR / "quickbyte_response.docx"
    doc.save(out)
    print(f"✓ Created {out}")


def _inr_words(n):
    """Very rough INR word form for realism."""
    if n >= 100000:
        return f"{n//100000} Lakh{' ' + str((n%100000)//1000) + ' Thousand' if (n%100000)//1000 else ''}"
    elif n >= 1000:
        return f"{n//1000} Thousand"
    return str(n)


# ════════════════════════════════════════════════════════════════════════════
# 4. DigitalEdge — Angled photo of rate card
# ════════════════════════════════════════════════════════════════════════════
def create_digitaledge_image():
    from PIL import Image, ImageDraw, ImageFont
    import math

    # Create a white "paper" rate card
    W, H = 1200, 2800
    img = Image.new("RGB", (W, H), color=(252, 252, 248))
    draw = ImageDraw.Draw(img)

    # Try to use a basic font; fall back to default
    try:
        font_title = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 36)
        font_h1    = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 26)
        font_body  = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 20)
        font_small = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 16)
    except Exception:
        font_title = font_h1 = font_body = font_small = ImageFont.load_default()

    # Header block
    draw.rectangle([0, 0, W, 100], fill=(0, 80, 160))
    draw.text((40, 20), "DigitalEdge Corp — Rate Card", font=font_title, fill=(255, 255, 255))
    draw.text((40, 65), "RFQ Ref: RFX-2024-ITH-001  |  Date: 25 Nov 2024  |  Valid 60 days", font=font_small, fill=(200, 220, 255))

    y = 120
    draw.text((40, y), "Contact: procurement@digitaledge.in | GST: 19AABCD5678D1ZB", font=font_small, fill=(80, 80, 80))
    y += 30
    draw.text((40, y), "ISO 9001:2021 Certified | Cert No: ISO9001-2021-DE-3309", font=font_small, fill=(0, 120, 0))
    y += 50

    # Column headers
    cols = [40, 120, 200, 580, 730, 870, 1020]
    headers = ["#", "Code", "Description", "Qty", "Unit", "Unit Price", "Total"]
    draw.rectangle([30, y, W-30, y+40], fill=(230, 240, 255))
    for cx, h in zip(cols, headers):
        draw.text((cx, y+10), h, font=font_h1, fill=(0, 60, 120))
    y += 45

    missing_de = {3, 16}
    # DigitalEdge edge case: CAB-001 quoted per 100m instead of per box (305m)
    row_colors = [(255, 255, 255), (245, 248, 255)]
    for i, item in enumerate(ITEMS):
        iid = item["id"]
        if iid in missing_de:
            continue
        price = PRICES[iid]["digitaledge"]
        draw.rectangle([30, y, W-30, y+36], fill=row_colors[i % 2])

        if price == "UNIT_UNCLEAR":
            # CAT6 cable — quoted per 100m, not per box
            unit_disp = "per 100m"
            price_disp = "₹2,680"   # ~₹8,200 / 3.05 (305m / 100m)
            total_disp = f"({item['qty']} × 305m boxes requested)"
            draw.text((cols[5], y+8), price_disp, font=font_body, fill=(180, 80, 0))
            draw.text((cols[6], y+8), total_disp, font=font_small, fill=(180, 80, 0))
        else:
            unit_disp = item["unit"]
            total = price * item["qty"]
            draw.text((cols[5], y+8), f"₹{price:,}", font=font_body, fill=(30, 30, 30))
            draw.text((cols[6], y+8), f"₹{total:,}", font=font_body, fill=(30, 30, 30))

        draw.text((cols[0], y+8), str(iid),              font=font_body, fill=(80, 80, 80))
        draw.text((cols[1], y+8), item["code"],          font=font_body, fill=(30, 30, 30))
        draw.text((cols[2], y+8), item["description"][:28], font=font_body, fill=(30, 30, 30))
        draw.text((cols[3], y+8), str(item["qty"]),      font=font_body, fill=(30, 30, 30))
        draw.text((cols[4], y+8), unit_disp,             font=font_body, fill=(30, 30, 30))
        y += 37

    # Footer note
    y += 20
    draw.text((40, y), "* CAT6 prices quoted per 100 linear metres. Buyer to confirm bundle size.", font=font_small, fill=(180, 80, 0))
    y += 25
    draw.text((40, y), "Items not quoted: SRV-003 (Blade Chassis), WLC-001 (Controller) — not in portfolio.", font=font_small, fill=(150, 0, 0))
    y += 25
    draw.text((40, y), "All prices in INR incl. duties, excl. GST. Warranty: 24 months all items. Lead time: 30 days.", font=font_small, fill=(60, 60, 60))

    # Crop to actual content
    img = img.crop((0, 0, W, min(y + 60, H)))

    # ── Simulate an angled phone photo ──────────────────────────────────────
    # 1. Slight rotation (phone wasn't held perfectly straight)
    angle = -4.5
    img_rot = img.rotate(angle, expand=True, fillcolor=(180, 175, 165))

    # 2. Add a "table surface" background (wood/desk texture — solid beige)
    pad = 80
    final_w = img_rot.width + pad * 2
    final_h = img_rot.height + pad * 2
    background = Image.new("RGB", (final_w, final_h), color=(185, 168, 140))
    background.paste(img_rot, (pad, pad))

    # 3. Perspective distortion (simulate phone angle) using a slight shear
    from PIL import ImageFilter
    # Add slight blur (camera shake / focus)
    background = background.filter(ImageFilter.GaussianBlur(radius=0.8))

    # 4. Slight brightness/contrast variation (uneven lighting)
    from PIL import ImageEnhance
    background = ImageEnhance.Brightness(background).enhance(0.92)
    background = ImageEnhance.Contrast(background).enhance(1.08)

    out = RESPONSES_DIR / "digitaledge_ratecard.png"
    background.save(out, "PNG", quality=88)
    print(f"✓ Created {out}")


# ════════════════════════════════════════════════════════════════════════════
# 5. Shree IT — Plain email
# ════════════════════════════════════════════════════════════════════════════
def create_shree_email():
    content = """From: shreeittrd@gmail.com
To: vishnu.vinod@democorp.in
Subject: Re: RFX-2024-ITH-001 — IT Hardware Quotation
Date: 26 November 2024, 11:38 AM

Dear Vishnu Sir,

Thanks for the RFQ. We are happy to quote for the items below. We have been
in business for 14 years and have supplied to many large companies.

Our rates (all prices in INR, per unit, excl GST):

SERVERS & COMPUTE:
- SRV-002 (Rack Server 1U): Rs.1,82,000/unit
- SRV-004 (Blade Module): Rs.89,000/unit
- WKS-001 (CAD Workstation): Not available with us, sorry
- DSK-001 (Desktop PC): Rs.34,500/unit

LAPTOPS:
- LPT-001 (High-end laptop): Rs.87,500/unit
- LPT-002 (Standard laptop): Rs.53,500/unit

NETWORKING:
- NSW-001 (24P Switch): Rs.16,800/unit
- NSW-002 (48P Switch): Rs.29,500/unit
- WAP-001 (WiFi AP): Rs.12,200/unit
- SFP-001 (SFP+ Module): Rs.3,200/unit

POWER:
- UPS-001 (1kVA): Rs.13,200/unit
- UPS-002 (3kVA): Rs.36,000/unit

STORAGE & COMPONENTS:
- SSD-001 (1TB SSD): Rs.16,500/unit
- MON-001 (Monitor 27"): Rs.30,000/unit
- RAK-001 (42U Rack): Rs.25,800/unit

CABLING:
- CAB-001 (CAT6 box 305m): Rs.7,500/box
- PAT-001 (Fiber Patch Panel): Rs.5,000/unit

For the following items, our rates are same as last year quote (you can check
our 2023 quotation ref SIT-2023-088):
- TNC-001 (Thin Client): same as last year
- PDU-001 (Rack PDU): same as last year
- RAM-001 (32GB DDR5): same as last year

Items SRV-001, SRV-003, RTR-001, FWL-001, NSW-003, WLC-001, KVM-001,
NAS-001, GPU-001 — we don't have these in stock currently. Can try to
arrange but cannot commit delivery within 30 days.

Freight: Extra. Bengaluru delivery free above Rs.5 lakh order value.
GST as applicable (18% on IT hardware).

Regarding questionnaire: We will provide ISO certificate if required
(we are in process of getting certified). GST number is 29AABCS9876E1ZC.
Payment Net 30 days. Warranty as per OEM.

Please call if any questions: +91-98765-43210 (Ramesh Shree, Proprietor)

Thanks and regards,
Ramesh Shree
Shree IT Traders, Bengaluru
"""
    out = RESPONSES_DIR / "shree_email.txt"
    out.write_text(content)
    print(f"✓ Created {out}")


# ════════════════════════════════════════════════════════════════════════════
# Main
# ════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    print("Creating vendor response files...")
    create_techpro_excel()

    try:
        create_globalit_pdf()
    except Exception as e:
        print(f"  PDF creation failed ({e}), creating text stub")
        _create_globalit_text_stub()

    create_quickbyte_docx()
    create_digitaledge_image()
    create_shree_email()
    print("\nAll vendor response files created in:", RESPONSES_DIR)
